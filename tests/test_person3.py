"""
Yakṣarāja — Person 3 Module Tests
========================================
Covers all Person 3 (LLM + CAM) modules:
  1. Research Agent       — run_research()
  2. Approval Agent       — write_bull_case()
  3. Dissent Agent        — write_bear_case()
  4. Coordinator          — synthesize_cam_recommendation()
  5. CAM Generator        — generate_cam()
  6. CEO Interview        — run_ceo_interview_analysis() fallback

All external API calls (Anthropic Claude, Tavily, Whisper) are mocked
so tests run offline, quickly, and deterministically.

Run:
    pytest tests/test_person3.py -v
    pytest tests/test_person3.py -v --tb=short
"""

import os
import sys
import json
import tempfile
import shutil
from unittest import mock
from unittest.mock import patch, MagicMock

import pytest

# ── Ensure project root is on sys.path ───────────────────────────────────────
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  SHARED FIXTURES                                                           ║
# ╚════════════════════════════════════════════════════════════════════════════╝

@pytest.fixture
def demo_company_data():
    """Standard synthetic company data dict used across multiple tests."""
    return {
        "company_name": "Sunrise Textile Mills",
        "sector": "Textiles",
        "fiscal_year": "FY2024",
        "revenue": 850.0,
        "ebitda": 127.5,
        "ebitda_margin": 0.15,
        "pat": 51.0,
        "net_margin": 0.06,
        "gross_margin": 0.38,
        "dscr": 1.85,
        "interest_coverage": 2.4,
        "debt_to_equity": 1.6,
        "current_ratio": 1.25,
        "roe": 0.14,
        "roa": 0.06,
        "cfo": 95.0,
        "free_cash_flow": 42.0,
        "total_debt": 520.0,
        "total_equity": 325.0,
        "total_assets": 1200.0,
        "beneish_m_score": -2.45,
        "altman_z_score": 2.3,
        "piotroski_f_score": 6,
        "auditor_distress_score": 1,
        "going_concern_flag": 0,
        "qualified_opinion_flag": 0,
        "auditor_resigned_flag": 0,
        "related_party_tx_to_rev": 0.05,
        "ensemble_pd": 0.12,
        "xgb_pd": 0.11,
        "rf_pd": 0.14,
        "lgb_pd": 0.13,
        "lending_decision": "APPROVE",
        "model_confidence": "HIGH_CONSENSUS",
        "model_disagreement_flag": "CONSENSUS",
        "risk_premium": 3.5,
        "satellite_activity_score": 82.5,
        "satellite_activity_category": "ACTIVE",
        "satellite_vs_revenue_flag": 0,
        "gst_vs_bank_divergence": 0.03,
        "gst_divergence_flag": 0,
        "gst_filing_delays_count": 1,
        "gst_payment_delay_days": 12,
        "contagion_risk_score": 0.15,
        "promoter_holding_pct": 0.62,
        "promoter_pledge_pct": 0.08,
        "promoter_npa_companies": 0,
        "promoter_total_companies": 4,
        "promoter_struck_off_companies": 0,
        "din_disqualified_count": 0,
        "network_npa_ratio": 0.0,
        "customer_concentration": 0.35,
        "supplier_concentration": 0.42,
        "ceo_sentiment_overall": 0.72,
        "ceo_specificity_score": 0.65,
        "ceo_deflection_score": 0.18,
        "ceo_overconfidence_score": 0.22,
        "ceo_sentiment_revenue": 0.70,
        "ceo_sentiment_debt": 0.45,
        "months_to_dscr_danger": 36,
        "label": 0,
    }


@pytest.fixture
def demo_research():
    """Standard research output dict."""
    return {
        "company_news_summary": "Sunrise Textile Mills reported 12% revenue growth in FY2024.",
        "industry_outlook": "POSITIVE",
        "research_sentiment_score": 0.72,
        "key_positives_found": [
            "PLI scheme support",
            "China+1 trend",
            "Domestic demand growth",
        ],
        "key_risks_found": [
            "Raw material price volatility",
            "Global demand slowdown risk",
        ],
        "promoter_red_flags": [],
        "research_sources": ["Economic Times", "Business Standard"],
        "used_fallback": False,
        "errors": [],
    }


@pytest.fixture
def demo_all_data(demo_company_data, demo_research):
    """Complete pipeline output dict for CAM generation."""
    return {
        "company_name": "Sunrise Textile Mills",
        "fiscal_year": 2024,
        "financial_data": demo_company_data,
        "forensics": {
            "beneish_m_score": -2.45,
            "altman_z_score": 2.3,
            "piotroski_f_score": 6,
        },
        "ml_scores": {
            "ensemble_pd": 0.12,
            "xgb_pd": 0.11,
            "rf_pd": 0.14,
            "lgb_pd": 0.13,
            "lending_decision": "CONDITIONAL_APPROVE",
            "risk_premium": 3.5,
        },
        "trajectory": {
            "dscr_trend": "STABLE",
            "months_to_danger": 36,
            "dscr_history": [1.55, 1.62, 1.70, 1.78, 1.85],
            "fiscal_years": [2020, 2021, 2022, 2023, 2024],
        },
        "network": {"contagion_risk_score": 0.15},
        "stress_test": {
            "dscr_p10": 1.05,
            "dscr_p50": 1.65,
            "dscr_p90": 2.15,
            "covenant_breach_probability": 0.08,
        },
        "dna_match": {
            "closest_default_archetype": "None (Healthy)",
            "max_archetype_similarity": 0.18,
        },
        "satellite": {
            "activity_score": 82.5,
            "activity_category": "ACTIVE",
            "vs_revenue_flag": 0,
        },
        "gst": {},
        "research": demo_research,
        "ceo_interview": {
            "key_scores": {
                "ceo_sentiment_overall": 0.45,
                "ceo_sentiment_debt": 0.15,
                "ceo_deflection_score": 0.18,
                "ceo_overconfidence_score": 0.12,
                "ceo_specificity_score": 0.55,
            },
            "red_flags": [],
            "red_flag_count": 0,
            "management_quality_score": 72.5,
            "used_fallback": False,
        },
        "bull_case": (
            "## 1. EXECUTIVE SUMMARY\n"
            "Strong lending opportunity with DSCR of 1.85x.\n\n"
            "## 2. FINANCIAL STRENGTHS\n"
            "- Healthy DSCR of 1.85x\n- Positive free cash flow\n"
        ),
        "bear_case": (
            "## 1. CRITICAL CONCERNS\n"
            "- D/E of 1.6x leaves limited buffer\n"
            "- Raw material volatility\n\n"
            "## 2. STRESS SCENARIO\n"
            "- Combined shock DSCR drops to 0.85x\n"
        ),
        "recommendation": {
            "lending_decision": "CONDITIONAL_APPROVE",
            "recommended_limit_cr": 187.0,
            "recommended_rate_pct": 10.0,
            "key_conditions": [
                "DSCR floor covenant at 1.20x",
                "Promoter personal guarantee",
                "Quarterly GST cross-verification",
                "Annual credit review",
            ],
            "bull_summary": "Strong DSCR and positive FCF.",
            "bear_summary": "Leverage risk and raw material exposure.",
            "final_rationale": "Conditional approve with ₹187 Cr at 10%.",
        },
    }


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  TEST 1 — RESEARCH AGENT                                                  ║
# ╚════════════════════════════════════════════════════════════════════════════╝

class TestResearchAgent:
    """Test research_agent.run_research() with mocked APIs → fallback path."""

    def test_research_agent_returns_structured_output(self):
        """
        Run research agent for 'Sunrise Textile Mills', sector 'Textiles'.
        With APIs mocked as unavailable, it should use the fallback database
        and still return a fully-structured dict.
        """
        # Mock the Tavily and Anthropic clients as unavailable so fallback is used
        with patch(
            "modules.person3_llm_cam.research_agent.TAVILY_AVAILABLE", False
        ), patch(
            "modules.person3_llm_cam.research_agent.ANTHROPIC_AVAILABLE", False
        ), patch(
            "modules.person3_llm_cam.research_agent.LANGGRAPH_AVAILABLE", False
        ):
            from modules.person3_llm_cam.research_agent import run_research

            result = run_research(
                company_name="Sunrise Textile Mills",
                sector="Textiles",
                promoter_name="Rajesh Kumar",
            )

        # ── Structure assertions ─────────────────────────────────────────
        assert isinstance(result, dict), "Result must be a dict"

        required_keys = [
            "company_news_summary",
            "industry_outlook",
            "key_risks_found",
            "key_positives_found",
            "promoter_red_flags",
            "research_sources",
            "research_sentiment_score",
            "used_fallback",
            "errors",
        ]
        for key in required_keys:
            assert key in result, f"Missing key: {key}"

        # ── Value assertions ─────────────────────────────────────────────
        assert isinstance(result["company_news_summary"], str)
        assert len(result["company_news_summary"]) > 0, "News summary should be non-empty"

        assert result["industry_outlook"] in ("POSITIVE", "NEUTRAL", "NEGATIVE"), \
            f"Invalid outlook: {result['industry_outlook']}"

        assert isinstance(result["key_risks_found"], list)
        assert isinstance(result["key_positives_found"], list)

        score = result["research_sentiment_score"]
        assert isinstance(score, (int, float)), "Sentiment score must be numeric"
        assert 0.0 <= score <= 1.0, f"Sentiment score {score} out of range [0, 1]"

    def test_research_agent_fallback_flag(self):
        """When APIs are unavailable, used_fallback should be True."""
        with patch(
            "modules.person3_llm_cam.research_agent.TAVILY_AVAILABLE", False
        ), patch(
            "modules.person3_llm_cam.research_agent.ANTHROPIC_AVAILABLE", False
        ), patch(
            "modules.person3_llm_cam.research_agent.LANGGRAPH_AVAILABLE", False
        ):
            from modules.person3_llm_cam.research_agent import run_research
            result = run_research("Sunrise Textile Mills", "Textiles")

        assert result["used_fallback"] is True, "Should flag fallback usage"


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  TEST 2 — APPROVAL AGENT (BULL CASE)                                      ║
# ╚════════════════════════════════════════════════════════════════════════════╝

class TestApprovalAgent:
    """Test approval_agent.write_bull_case() with mocked Claude → fallback."""

    def test_approval_agent_produces_bull_case(self, demo_company_data, demo_research):
        """
        Run approval agent with synthetic data.
        Should produce a substantial bull case text mentioning financial ratios.
        """
        with patch(
            "modules.person3_llm_cam.approval_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.approval_agent import write_bull_case

            result = write_bull_case(demo_company_data, demo_research)

        # ── Length assertions ────────────────────────────────────────────
        assert isinstance(result, str), "Bull case must be a string"
        assert len(result) > 200, \
            f"Bull case too short ({len(result)} chars), expected > 200"

        # ── Content assertions ───────────────────────────────────────────
        # Must mention at least one financial ratio keyword
        ratio_keywords = [
            "DSCR", "dscr", "debt", "margin", "coverage", "ROE", "roe",
            "interest", "equity", "EBITDA", "ebitda", "revenue",
        ]
        mentions_ratio = any(kw in result for kw in ratio_keywords)
        assert mentions_ratio, \
            "Bull case must mention at least one financial ratio/metric"

    def test_bull_case_is_non_empty_with_sections(self, demo_company_data, demo_research):
        """Fallback bull case should contain structured sections."""
        with patch(
            "modules.person3_llm_cam.approval_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.approval_agent import write_bull_case
            result = write_bull_case(demo_company_data, demo_research)

        # Should have multiple paragraphs / sections
        lines = [l.strip() for l in result.split("\n") if l.strip()]
        assert len(lines) >= 5, \
            f"Bull case should have at least 5 non-empty lines, got {len(lines)}"


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  TEST 3 — DISSENT AGENT (BEAR CASE)                                       ║
# ╚════════════════════════════════════════════════════════════════════════════╝

class TestDissentAgent:
    """Test dissent_agent.write_bear_case() with mocked Claude → fallback."""

    def test_dissent_agent_produces_bear_case(self, demo_company_data, demo_research):
        """
        Run dissent agent on same data + approval output.
        Should produce counter-arguments that differ from the bull case.
        """
        # First generate a bull case to feed into the dissent agent
        with patch(
            "modules.person3_llm_cam.approval_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.approval_agent import write_bull_case
            bull_case = write_bull_case(demo_company_data, demo_research)

        # Now generate bear case
        with patch(
            "modules.person3_llm_cam.dissent_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.dissent_agent import write_bear_case
            bear_case = write_bear_case(demo_company_data, bull_case, demo_research)

        # ── Structure assertions ─────────────────────────────────────────
        assert isinstance(bear_case, str), "Bear case must be a string"
        assert len(bear_case) > 200, \
            f"Bear case too short ({len(bear_case)} chars), expected > 200"

        # ── Counter-argument count ───────────────────────────────────────
        # Count substantive lines (bullets, numbered items, or sentences with "risk"/"concern")
        counter_indicators = ["-", "•", "concern", "risk", "weakness",
                              "challenge", "flag", "danger", "vulnerable"]
        counter_count = 0
        for line in bear_case.split("\n"):
            line_lower = line.strip().lower()
            if any(ind in line_lower for ind in counter_indicators) and len(line.strip()) > 20:
                counter_count += 1
        assert counter_count >= 3, \
            f"Bear case should have at least 3 counter-arguments, found {counter_count}"

        # ── Different from bull case ─────────────────────────────────────
        assert bear_case != bull_case, "Bear case must differ from bull case"

        # Check they don't share the same first 100 chars (structurally different)
        assert bear_case[:100] != bull_case[:100], \
            "Bear case opening should differ from bull case opening"


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  TEST 4 — COORDINATOR (RECOMMENDATION SYNTHESIS)                          ║
# ╚════════════════════════════════════════════════════════════════════════════╝

class TestCoordinator:
    """Test synthesize_cam_recommendation() with mocked Claude → fallback."""

    def test_coordinator_produces_valid_recommendation(self, demo_company_data):
        """
        Run coordinator on bull + bear cases.
        Should produce a valid recommendation dict with correct structure and ranges.
        """
        bull_case = (
            "## BULL CASE\n"
            "Strong DSCR of 1.85x supports lending. "
            "Positive free cash flow of ₹42 Cr. "
            "PLI scheme provides margin cushion."
        )
        bear_case = (
            "## BEAR CASE\n"
            "D/E of 1.6x limits buffer. "
            "Raw material price volatility threatens margins. "
            "Customer concentration at 35%."
        )

        scores = {
            "ensemble_pd": 0.12,
            "dscr": 1.85,
            "lending_decision": "APPROVE",
            "risk_premium": 3.5,
            "revenue": 850.0,
            "promoter_pledge_pct": 0.08,
            "contagion_risk_score": 0.15,
        }

        with patch(
            "modules.person3_llm_cam.dissent_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.dissent_agent import synthesize_cam_recommendation
            result = synthesize_cam_recommendation(bull_case, bear_case, scores)

        # ── Structure assertions ─────────────────────────────────────────
        assert isinstance(result, dict), "Recommendation must be a dict"

        required_keys = [
            "lending_decision",
            "recommended_limit_cr",
            "recommended_rate_pct",
            "key_conditions",
            "bull_summary",
            "bear_summary",
            "final_rationale",
        ]
        for key in required_keys:
            assert key in result, f"Missing key: {key}"

        # ── Value assertions ─────────────────────────────────────────────
        valid_decisions = ["APPROVE", "CONDITIONAL_APPROVE", "REJECT"]
        assert result["lending_decision"] in valid_decisions, \
            f"Invalid decision: {result['lending_decision']}, expected one of {valid_decisions}"

        rate = result["recommended_rate_pct"]
        assert isinstance(rate, (int, float)), "Rate must be numeric"
        assert 8.0 <= rate <= 20.0, \
            f"Rate {rate}% out of expected range [8.0, 20.0]"

        limit = result["recommended_limit_cr"]
        assert isinstance(limit, (int, float)), "Limit must be numeric"
        assert limit > 0, "Limit must be positive"

        assert isinstance(result["key_conditions"], list), "Conditions must be a list"
        assert len(result["key_conditions"]) >= 1, "Should have at least 1 condition"

        assert isinstance(result["final_rationale"], str)
        assert len(result["final_rationale"]) > 20, "Rationale should be substantive"

    def test_coordinator_reject_for_high_pd(self):
        """High PD + low DSCR should trigger REJECT decision."""
        scores_bad = {
            "ensemble_pd": 0.65,
            "dscr": 0.8,
            "lending_decision": "REJECT",
            "risk_premium": 8.0,
            "revenue": 200.0,
        }

        with patch(
            "modules.person3_llm_cam.dissent_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.dissent_agent import synthesize_cam_recommendation
            result = synthesize_cam_recommendation("bull", "bear", scores_bad)

        assert result["lending_decision"] == "REJECT", \
            f"Expected REJECT for PD=0.65, DSCR=0.8, got {result['lending_decision']}"

    def test_coordinator_approve_for_strong_company(self):
        """Low PD + high DSCR should trigger APPROVE decision."""
        scores_good = {
            "ensemble_pd": 0.08,
            "dscr": 2.5,
            "lending_decision": "APPROVE",
            "risk_premium": 2.0,
            "revenue": 1000.0,
        }

        with patch(
            "modules.person3_llm_cam.dissent_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.dissent_agent import synthesize_cam_recommendation
            result = synthesize_cam_recommendation("bull", "bear", scores_good)

        assert result["lending_decision"] == "APPROVE", \
            f"Expected APPROVE for PD=0.08, DSCR=2.5, got {result['lending_decision']}"


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  TEST 5 — CAM GENERATOR                                                   ║
# ╚════════════════════════════════════════════════════════════════════════════╝

class TestCAMGenerator:
    """Test cam_generator.generate_cam() creates valid DOCX files."""

    def test_cam_generator_creates_docx(self, demo_all_data):
        """
        Run generate_cam with synthetic data.
        Should produce a DOCX file > 10KB that python-docx can open.
        """
        # Use a temp directory for output
        output_dir = tempfile.mkdtemp(prefix="intelli_credit_test_")

        try:
            from modules.person3_llm_cam.cam_generator import generate_cam

            result_path = generate_cam(demo_all_data, output_dir=output_dir)

            # ── File existence ───────────────────────────────────────────
            assert result_path, "generate_cam should return a non-empty path"
            assert os.path.exists(result_path), f"DOCX file not found: {result_path}"

            # ── File size ────────────────────────────────────────────────
            file_size = os.path.getsize(result_path)
            assert file_size > 10 * 1024, \
                f"DOCX too small ({file_size} bytes), expected > 10KB"

            # ── Valid DOCX (openable by python-docx) ─────────────────────
            from docx import Document
            doc = Document(result_path)
            assert len(doc.paragraphs) > 10, \
                f"DOCX should have many paragraphs, found {len(doc.paragraphs)}"

            # ── Contains expected content ────────────────────────────────
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Sunrise Textile Mills" in full_text, \
                "DOCX should contain company name"
            assert "CREDIT APPRAISAL MEMORANDUM" in full_text, \
                "DOCX should contain the title"

            # ── JSON scores file also created ────────────────────────────
            json_files = [f for f in os.listdir(output_dir) if f.endswith(".json")]
            assert len(json_files) >= 1, "Should also create a scores JSON file"

            json_path = os.path.join(output_dir, json_files[0])
            with open(json_path) as f:
                scores = json.load(f)
            assert "company_name" in scores
            assert "lending_decision" in scores

        finally:
            # Cleanup temp directory
            shutil.rmtree(output_dir, ignore_errors=True)

    def test_cam_generator_handles_missing_data(self):
        """CAM generator should handle minimal/missing data without crashing."""
        minimal_data = {
            "company_name": "Test Corp",
            "fiscal_year": 2024,
            "financial_data": {"revenue": 100.0, "dscr": 1.2},
            # All other keys missing → should show N/A gracefully
        }
        output_dir = tempfile.mkdtemp(prefix="intelli_credit_test_")
        try:
            from modules.person3_llm_cam.cam_generator import generate_cam
            result_path = generate_cam(minimal_data, output_dir=output_dir)

            assert result_path, "Should succeed even with minimal data"
            assert os.path.exists(result_path), "DOCX should be created"
        finally:
            shutil.rmtree(output_dir, ignore_errors=True)

    def test_cam_filename_format(self, demo_all_data):
        """DOCX filename should follow CAM_[company]_[date].docx pattern."""
        output_dir = tempfile.mkdtemp(prefix="intelli_credit_test_")
        try:
            from modules.person3_llm_cam.cam_generator import generate_cam
            result_path = generate_cam(demo_all_data, output_dir=output_dir)

            filename = os.path.basename(result_path)
            assert filename.startswith("CAM_"), \
                f"Filename should start with 'CAM_', got: {filename}"
            assert filename.endswith(".docx"), \
                f"Filename should end with '.docx', got: {filename}"
            assert "Sunrise_Textile_Mills" in filename, \
                f"Filename should contain company name, got: {filename}"
        finally:
            shutil.rmtree(output_dir, ignore_errors=True)


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  TEST 6 — CEO INTERVIEW FALLBACK                                          ║
# ╚════════════════════════════════════════════════════════════════════════════╝

class TestCEOInterviewFallback:
    """Test ceo_interview.run_ceo_interview_analysis() fallback mode."""

    def test_ceo_interview_fallback_returns_valid_dict(self):
        """
        Run CEO interview module with no audio file and no transcript.
        Should return a valid dict with all required keys via fallback.
        """
        company_data = {
            "dscr": 1.85,
            "debt_to_equity": 1.6,
            "net_margin": 0.06,
            "label": 0,
        }

        from modules.person3_llm_cam.ceo_interview import run_ceo_interview_analysis

        result = run_ceo_interview_analysis(
            audio_path=None,
            transcript=None,
            company_data=company_data,
        )

        # ── Structure assertions ─────────────────────────────────────────
        assert isinstance(result, dict), "Result must be a dict"

        required_keys = [
            "key_scores",
            "red_flags",
            "red_flag_count",
            "management_quality_score",
            "used_fallback",
        ]
        for key in required_keys:
            assert key in result, f"Missing key: {key}"

        # ── Key scores sub-dict ──────────────────────────────────────────
        score_keys = [
            "ceo_sentiment_overall",
            "ceo_sentiment_debt",
            "ceo_deflection_score",
            "ceo_overconfidence_score",
            "ceo_specificity_score",
        ]
        for key in score_keys:
            assert key in result["key_scores"], f"Missing score key: {key}"
            val = result["key_scores"][key]
            assert isinstance(val, (int, float)), f"{key} must be numeric, got {type(val)}"
            assert 0.0 <= val <= 1.0, f"{key} = {val} out of range [0, 1]"

        # ── MQ score range ───────────────────────────────────────────────
        mq = result["management_quality_score"]
        assert isinstance(mq, (int, float)), "MQ score must be numeric"
        assert 0.0 <= mq <= 100.0, \
            f"MQ score {mq} out of range [0, 100]"

        # ── Fallback flag ────────────────────────────────────────────────
        assert result["used_fallback"] is True, "Should flag as fallback"

    def test_ceo_interview_fallback_no_company_data(self):
        """Fallback should work even with no company data (None)."""
        from modules.person3_llm_cam.ceo_interview import run_ceo_interview_analysis

        result = run_ceo_interview_analysis(
            audio_path=None,
            transcript=None,
            company_data=None,
        )

        assert isinstance(result, dict)
        assert "management_quality_score" in result
        assert 0.0 <= result["management_quality_score"] <= 100.0

    def test_ceo_interview_healthy_vs_stressed_company(self):
        """
        A healthy company (high DSCR, low D/E) should get better MQ scores
        than a stressed company (low DSCR, high D/E, label=1).
        """
        from modules.person3_llm_cam.ceo_interview import run_ceo_interview_analysis

        healthy = run_ceo_interview_analysis(
            company_data={"dscr": 2.5, "debt_to_equity": 0.8, "net_margin": 0.12, "label": 0}
        )
        stressed = run_ceo_interview_analysis(
            company_data={"dscr": 0.7, "debt_to_equity": 3.5, "net_margin": -0.02, "label": 1}
        )

        assert healthy["management_quality_score"] > stressed["management_quality_score"], \
            f"Healthy MQ ({healthy['management_quality_score']}) should exceed " \
            f"stressed MQ ({stressed['management_quality_score']})"

    def test_management_quality_score_function(self):
        """Test get_management_quality_score() directly with known inputs."""
        from modules.person3_llm_cam.ceo_interview import get_management_quality_score

        # Perfect scores → should be near 100
        perfect = {
            "key_scores": {
                "ceo_sentiment_overall": 0.4,   # balanced
                "ceo_sentiment_debt": -0.1,      # honestly negative about debt
                "ceo_deflection_score": 0.0,     # no deflection
                "ceo_overconfidence_score": 0.1,  # mild confidence
                "ceo_specificity_score": 0.6,     # very specific
            },
            "red_flags": [],
        }
        score_perfect = get_management_quality_score(perfect)
        assert score_perfect >= 80.0, \
            f"Perfect inputs should score ≥80, got {score_perfect}"

        # Terrible scores → should be low
        terrible = {
            "key_scores": {
                "ceo_sentiment_overall": 0.9,     # suspiciously positive
                "ceo_sentiment_debt": 0.8,         # suspiciously positive about debt
                "ceo_deflection_score": 0.8,       # highly evasive
                "ceo_overconfidence_score": 0.6,   # extreme overconfidence
                "ceo_specificity_score": 0.05,     # almost no specifics
            },
            "red_flags": [
                {"flag_type": "a", "severity": "HIGH", "description": "x"},
                {"flag_type": "b", "severity": "HIGH", "description": "y"},
                {"flag_type": "c", "severity": "HIGH", "description": "z"},
            ],
        }
        score_terrible = get_management_quality_score(terrible)
        assert score_terrible < 30.0, \
            f"Terrible inputs should score <30, got {score_terrible}"


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  TEST — INTEGRATION: FULL ADVERSARIAL PIPELINE (fallback mode)            ║
# ╚════════════════════════════════════════════════════════════════════════════╝

class TestAdversarialPipelineIntegration:
    """Integration test: bull → bear → coordinator, all in fallback mode."""

    def test_full_adversarial_pipeline(self, demo_company_data, demo_research):
        """
        Run the full adversarial pipeline end-to-end in fallback mode:
        write_bull_case → write_bear_case → synthesize_cam_recommendation
        """
        with patch(
            "modules.person3_llm_cam.approval_agent.ANTHROPIC_AVAILABLE", False
        ), patch(
            "modules.person3_llm_cam.dissent_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.approval_agent import write_bull_case
            from modules.person3_llm_cam.dissent_agent import (
                write_bear_case,
                synthesize_cam_recommendation,
            )

            # Step 1: Bull case
            bull = write_bull_case(demo_company_data, demo_research)
            assert isinstance(bull, str) and len(bull) > 100

            # Step 2: Bear case (reads the bull case)
            bear = write_bear_case(demo_company_data, bull, demo_research)
            assert isinstance(bear, str) and len(bear) > 100
            assert bear != bull

            # Step 3: Coordinator
            rec = synthesize_cam_recommendation(bull, bear, demo_company_data)
            assert rec["lending_decision"] in [
                "APPROVE", "CONDITIONAL_APPROVE", "REJECT"
            ]
            assert isinstance(rec["recommended_limit_cr"], (int, float))
            assert isinstance(rec["recommended_rate_pct"], (int, float))
            assert len(rec["key_conditions"]) >= 1


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  TEST — CORRECTED CAM RULES (Rules 1-7)                                   ║
# ╚════════════════════════════════════════════════════════════════════════════╝

class TestCorrectedRules:
    """Verify the corrected system prompt rules are enforced in code."""

    def test_dscr_below_1_forces_reject(self):
        """Rule 4a: DSCR < 1.0 → coordinator MUST return REJECT with limit=0."""
        scores_dscr_below_1 = {
            "ensemble_pd": 0.15,
            "dscr": 0.9,
            "lending_decision": "CONDITIONAL_APPROVE",
            "risk_premium": 3.5,
            "revenue": 850.0,
        }

        with patch(
            "modules.person3_llm_cam.dissent_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.dissent_agent import synthesize_cam_recommendation
            result = synthesize_cam_recommendation("bull", "bear", scores_dscr_below_1)

        assert result["lending_decision"] == "REJECT", \
            f"Expected REJECT for DSCR=0.9, got {result['lending_decision']}"
        assert result["recommended_limit_cr"] == 0.0, \
            f"Expected limit=0 for REJECT, got {result['recommended_limit_cr']}"

    def test_pd_above_40_forces_reject(self):
        """Rule 4e: PD > 0.40 → coordinator MUST return REJECT."""
        scores_high_pd = {
            "ensemble_pd": 0.45,
            "dscr": 1.5,
            "lending_decision": "CONDITIONAL_APPROVE",
            "risk_premium": 5.0,
            "revenue": 500.0,
        }

        with patch(
            "modules.person3_llm_cam.dissent_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.dissent_agent import synthesize_cam_recommendation
            result = synthesize_cam_recommendation("bull", "bear", scores_high_pd)

        assert result["lending_decision"] == "REJECT", \
            f"Expected REJECT for PD=45%, got {result['lending_decision']}"
        assert result["recommended_limit_cr"] == 0.0, \
            f"Expected limit=0 for REJECT, got {result['recommended_limit_cr']}"

    def test_missing_data_shows_data_required(self):
        """Rule 1: Missing financial figures → [DATA REQUIRED] in bull case."""
        from modules.person3_llm_cam.approval_agent import _display_val, _display_pct, _display_ratio

        assert _display_val(None) == "[DATA REQUIRED]"
        assert _display_val(0.0) == "[DATA REQUIRED]"
        assert _display_pct(None) == "[DATA REQUIRED]"
        assert _display_ratio(None) == "[DATA REQUIRED]"

        # Valid values should NOT show [DATA REQUIRED]
        assert "[DATA REQUIRED]" not in _display_val(850.0)
        assert "[DATA REQUIRED]" not in _display_pct(0.15)
        assert "[DATA REQUIRED]" not in _display_ratio(1.85)

    def test_sector_appropriate_stress_in_bear_case(self):
        """Rule 3: Bear case stress scenarios must match company sector."""
        steel_company = {
            "company_name": "Tata Steel",
            "sector": "Steel",
            "fiscal_year": "FY2024",
            "revenue": 2000.0, "ebitda": 400.0, "ebitda_margin": 0.20,
            "pat": 200.0, "net_margin": 0.10,
            "dscr": 1.8, "interest_coverage": 3.0,
            "debt_to_equity": 1.2, "current_ratio": 1.3,
            "roe": 0.15, "roa": 0.08,
            "cfo": 350.0, "free_cash_flow": 150.0,
            "total_debt": 800.0, "total_equity": 667.0,
            "beneish_m_score": -2.8, "altman_z_score": 2.5,
            "piotroski_f_score": 7, "auditor_distress_score": 0,
            "going_concern_flag": 0, "qualified_opinion_flag": 0,
            "related_party_tx_to_rev": 0.03,
            "ensemble_pd": 0.10, "model_disagreement_flag": "CONSENSUS",
            "risk_premium": 3.0,
            "satellite_activity_score": 90.0,
            "satellite_activity_category": "ACTIVE",
            "satellite_vs_revenue_flag": 0,
            "gst_vs_bank_divergence": 0.02, "gst_divergence_flag": 0,
            "gst_filing_delays_count": 0, "gst_payment_delay_days": 5,
            "contagion_risk_score": 0.10,
            "promoter_pledge_pct": 0.05,
            "promoter_npa_companies": 0, "din_disqualified_count": 0,
            "ceo_deflection_score": 0.15,
            "ceo_overconfidence_score": 0.20,
            "ceo_specificity_score": 0.60,
            "ceo_sentiment_revenue": 0.65,
            "ceo_sentiment_debt": 0.25,
        }
        steel_research = {
            "key_risks_found": ["Iron ore price volatility"],
            "key_positives_found": ["Infrastructure demand growth"],
            "industry_outlook": "POSITIVE",
            "research_sentiment_score": 0.70,
        }

        with patch(
            "modules.person3_llm_cam.dissent_agent.ANTHROPIC_AVAILABLE", False
        ):
            from modules.person3_llm_cam.dissent_agent import write_bear_case
            bear = write_bear_case(steel_company, "Bull case text", steel_research)

        # Steel company bear case should mention steel-specific stresses
        bear_lower = bear.lower()
        steel_keywords = ["iron ore", "coking coal", "cbam", "steel", "china"]
        mentions_steel = any(kw in bear_lower for kw in steel_keywords)
        assert mentions_steel, \
            f"Steel company bear case should mention steel-specific stresses, but didn't. Got:\n{bear[:500]}"

        # Should NOT mention cotton/textile stresses
        assert "cotton" not in bear_lower, \
            "Steel company bear case should NOT mention cotton"

    def test_cam_validation_forces_reject_on_dscr(self):
        """Rule 4a: validate_cam_consistency forces REJECT when DSCR < 1.0."""
        from modules.person3_llm_cam.cam_generator import validate_cam_consistency

        data = {
            "financial_data": {"dscr": 0.8, "cfo": -10.0},
            "ml_scores": {"ensemble_pd": 0.15},
            "recommendation": {
                "lending_decision": "CONDITIONAL_APPROVE",
                "recommended_limit_cr": 100.0,
            },
            "ceo_interview": {},
        }

        issues = validate_cam_consistency(data)
        assert any(i["rule"] == "4a" for i in issues), \
            "Should flag DSCR < 1.0 as rule 4a violation"
        assert data["recommendation"]["lending_decision"] == "REJECT", \
            "Should force REJECT when DSCR < 1.0"
        assert data["recommendation"]["recommended_limit_cr"] == 0.0, \
            "Should force limit to 0 when REJECT"

    def test_cam_validation_overrides_archetype_for_positive_cfo(self):
        """Rule 6: Positive CFO → override default archetype to 'None (Healthy)'."""
        from modules.person3_llm_cam.cam_generator import validate_cam_consistency

        data = {
            "financial_data": {"dscr": 2.0, "cfo": 100.0},
            "ml_scores": {"ensemble_pd": 0.10},
            "recommendation": {"lending_decision": "APPROVE"},
            "dna_match": {
                "closest_default_archetype": "Leveraged Buyout Failure",
                "max_archetype_similarity": 0.30,
            },
            "ceo_interview": {},
        }

        issues = validate_cam_consistency(data)
        assert any(i["rule"] == "6" for i in issues), \
            "Should flag mismatched archetype for positive CFO"
        assert data["dna_match"]["closest_default_archetype"] == "None (Healthy)", \
            "Should override archetype to 'None (Healthy)' for positive CFO"


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  MAIN — allow running with `python tests/test_person3.py`                 ║
# ╚════════════════════════════════════════════════════════════════════════════╝

if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
