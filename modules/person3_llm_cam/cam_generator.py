"""
Yakṣarāja — CAM Document Generator (Person 3)
=====================================================
Assembles ALL outputs from Person 1 (ML Core) + Person 2 (Alt Data) +
Person 3 (LLM Agents) into a professional Credit Appraisal Memo (CAM)
as a DOCX file.

Parts:
  A — Data assembly from pipeline
  B — 11 section-generator functions
  C — DOCX assembly with professional styling
  D — Save DOCX + summary JSON

Author: Person 3
Module: modules/person3_llm_cam/cam_generator.py
"""

import os
import re
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional

from loguru import logger

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not installed — CAM generation will be unavailable")
    DOCX_AVAILABLE = False


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  CONSTANTS & COLORS                                                       •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

# Vivriti Capital inspired palette
NAVY        = RGBColor(0x0A, 0x1F, 0x3C) if DOCX_AVAILABLE else None
DARK_BLUE   = RGBColor(0x14, 0x3A, 0x6B) if DOCX_AVAILABLE else None
ORANGE      = RGBColor(0xE8, 0x6C, 0x00) if DOCX_AVAILABLE else None
GREEN       = RGBColor(0x1B, 0x7A, 0x2B) if DOCX_AVAILABLE else None
RED         = RGBColor(0xC6, 0x28, 0x28) if DOCX_AVAILABLE else None
AMBER       = RGBColor(0xE6, 0x8A, 0x00) if DOCX_AVAILABLE else None
GREY        = RGBColor(0x75, 0x75, 0x75) if DOCX_AVAILABLE else None
WHITE       = RGBColor(0xFF, 0xFF, 0xFF) if DOCX_AVAILABLE else None
BLACK       = RGBColor(0x00, 0x00, 0x00) if DOCX_AVAILABLE else None
LIGHT_GREY  = RGBColor(0xF0, 0xF0, 0xF0) if DOCX_AVAILABLE else None

HEADER_BG   = "0A1F3C"  # hex for XML shading
ROW_ALT_BG  = "F5F5F5"


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  HELPER: SAFE DATA ACCESS                                                •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def _g(data: dict, *keys, default="Not Given"):
    """Safely get nested values: _g(data, 'ml_scores', 'ensemble_pd', default=0)."""
    current = data
    for key in keys:
        if isinstance(current, dict):
            current = current.get(key, default)
        else:
            return default
    return current if current is not None else default


def _fmt(value, fmt=".2f", suffix="", prefix=""):
    """Format a numeric value safely."""
    if value is None or value in ("N/A", "Not Given"):
        return "Not Given"
    try:
        return f"{prefix}{value:{fmt}}{suffix}"
    except (ValueError, TypeError):
        return str(value)


def _pct(value, decimals=1):
    """Format as percentage."""
    if value is None or value in ("N/A", "Not Given"):
        return "Not Given"
    try:
        return f"{float(value) * 100:.{decimals}f}%"
    except (ValueError, TypeError):
        return str(value)


def _not_given(val) -> bool:
    """Return True if the value represents missing/not-given data."""
    return val is None or val in ("N/A", "Not Given", "")


def _risk_color(level: str):
    """Return RGBColor for risk level string."""
    level = str(level).upper()
    if level in ("HIGH", "RED", "REJECT", "NEGATIVE"):
        return RED
    elif level in ("MEDIUM", "AMBER", "CONDITIONAL_APPROVE", "REVIEW", "NEUTRAL"):
        return AMBER
    else:
        return GREEN


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  CONSISTENCY VALIDATION (Rules 4, 5, 6, 7)                                •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def validate_cam_consistency(data: dict) -> list:
    """
    Validate CAM data for internal consistency before document assembly.
    Returns a list of warning/error dicts: {"rule": str, "level": str, "message": str}.
    Also mutates data to enforce mandatory corrections (e.g., DSCR < 1.0 → REJECT).
    """
    issues = []
    fin = data.get("financial_data", {})
    rec = data.get("recommendation", {})
    ml = data.get("ml_scores", {})
    ceo = data.get("ceo_interview", {})

    dscr = fin.get("dscr", rec.get("dscr", ml.get("dscr")))
    pd_val = ml.get("ensemble_pd", fin.get("ensemble_pd"))
    decision = rec.get("lending_decision", ml.get("lending_decision", ""))
    cfo = fin.get("cfo")

    # Rule 4a: DSCR < 1.0 → must be REJECT
    if dscr is not None and float(dscr) < 1.0 and decision != "REJECT":
        issues.append({
            "rule": "4a",
            "level": "CRITICAL",
            "message": f"DSCR {dscr:.2f}x < 1.0 but decision is {decision} — forcing REJECT",
        })
        if isinstance(rec, dict):
            rec["lending_decision"] = "REJECT"
            rec["recommended_limit_cr"] = 0.0

    # Rule 4c: Negative gross margin + positive EBITDA margin = inconsistent
    gross_margin = fin.get("gross_margin")
    ebitda_margin = fin.get("ebitda_margin")
    if gross_margin is not None and ebitda_margin is not None:
        if float(gross_margin) < 0 and float(ebitda_margin) > 0:
            issues.append({
                "rule": "4c",
                "level": "WARNING",
                "message": f"Gross margin {gross_margin:.1%} is negative but EBITDA margin {ebitda_margin:.1%} is positive — check data",
            })

    # Rule 4e: PD > 40% → must be REJECT
    if pd_val is not None and float(pd_val) > 0.40 and decision != "REJECT":
        issues.append({
            "rule": "4e",
            "level": "CRITICAL",
            "message": f"PD {pd_val:.2%} > 40% but decision is {decision} — forcing REJECT",
        })
        if isinstance(rec, dict):
            rec["lending_decision"] = "REJECT"
            rec["recommended_limit_cr"] = 0.0

    # Rule 6: Strong positive CFO → no close default archetype
    dna = data.get("dna_match", {})
    if cfo is not None and float(cfo) > 0:
        closest = dna.get("closest_default_archetype", "")
        if closest and closest not in ("None (Healthy)", "N/A", ""):
            issues.append({
                "rule": "6",
                "level": "WARNING",
                "message": f"CFO is positive (₹{cfo:.1f} Cr) but default archetype is '{closest}' — overriding to 'None (Healthy)'",
            })
            dna["closest_default_archetype"] = "None (Healthy)"

    # Rule 7: Promoter holding should not be a suspicious default
    promoter_pct = fin.get("promoter_holding_pct")
    if promoter_pct is None:
        issues.append({
            "rule": "7",
            "level": "INFO",
            "message": "Promoter holding percentage not provided — will show as N/A",
        })

    # Log all issues
    for issue in issues:
        log_fn = logger.error if issue["level"] == "CRITICAL" else logger.warning
        log_fn(f"CAM Validation [{issue['rule']}]: {issue['message']}")

    return issues


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  DOCX STYLING HELPERS                                                     •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_styled_heading(doc, text: str, level: int = 1):
    """Add a heading with custom navy color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = NAVY if level == 1 else DARK_BLUE
    return heading


def _add_para(doc, text: str, bold: bool = False, italic: bool = False,
              color=None, size: int = 10, alignment=None, space_after: int = 6):
    """Add a styled paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if alignment:
        para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    return para


def _add_md_inline_para(doc, text: str, base_size: int = 10, color=None, space_after: int = 4):
    """Add a paragraph with **bold** inline markdown rendered as actual bold runs."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.size = Pt(base_size)
        run.bold = (i % 2 == 1)  # odd indices are inside **
        if color:
            run.font.color.rgb = color
    return para


def _add_md_line(doc, line: str, heading_color=None, text_color=None, base_size: int = 10):
    """
    Parse one markdown-formatted line and add it as a styled DOCX paragraph.
    Handles: # headings, * / - bullets, **inline bold**, --- separators.
    """
    line = line.strip()
    if not line or re.match(r'^-{3,}$', line) or re.match(r'^={3,}$', line):
        return  # Skip empty lines and horizontal rules

    # Detect heading level (# / ## / ### / ####)
    hm = re.match(r'^(#{1,4})\s+(.*)', line)
    if hm:
        level = len(hm.group(1))
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', hm.group(2)).strip()
        text = re.sub(r'\*(.*?)\*', r'\1', text).strip()
        sizes = {1: base_size + 4, 2: base_size + 2, 3: base_size + 1, 4: base_size}
        _add_para(doc, text, bold=True, size=sizes.get(level, base_size),
                  color=heading_color, space_after=3)
        return

    # Detect bullet point (* or - or •, optionally indented)
    bm = re.match(r'^[ \t]*[*\-•]\s+(.*)', line)
    if bm:
        _add_md_inline_para(doc, f"  \u2022 {bm.group(1).strip()}",
                            base_size=base_size - 1, color=text_color, space_after=2)
        return

    # Regular paragraph with possible inline bold
    _add_md_inline_para(doc, line, base_size=base_size, color=text_color, space_after=4)


def _add_key_value(doc, key: str, value: str, color=None):
    """Add a key: value line with bold key."""
    para = doc.add_paragraph()
    run_key = para.add_run(f"{key}: ")
    run_key.bold = True
    run_key.font.size = Pt(10)
    run_val = para.add_run(str(value))
    run_val.font.size = Pt(10)
    if color:
        run_val.font.color.rgb = color
    para.paragraph_format.space_after = Pt(3)
    return para


def _create_table(doc, headers: list, rows: list, col_widths: list = None):
    """
    Create a styled table with header row and alternating row shading.
    headers: list of column header strings
    rows: list of lists (each inner list = one row of values)
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val) if val is not None else "Not Given"
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx % 2 == 1:
                _set_cell_shading(cell, ROW_ALT_BG)

    # Column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def _add_risk_badge(doc, label: str, level: str):
    """Add a colored risk indicator line: — LABEL: LEVEL."""
    color = _risk_color(level)
    para = doc.add_paragraph()
    indicator = para.add_run("— ")
    indicator.font.color.rgb = color
    indicator.font.size = Pt(11)
    indicator.bold = True
    key_run = para.add_run(f"{label}: ")
    key_run.bold = True
    key_run.font.size = Pt(10)
    val_run = para.add_run(str(level))
    val_run.font.color.rgb = color
    val_run.bold = True
    val_run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(2)


def _add_page_break(doc):
    """Insert a page break."""
    doc.add_page_break()


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 1 — COVER PAGE                                                  •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_cover_page(doc, data: dict):
    """Generate a polished CAM cover page with header banner, key metrics table."""
    company_name = str(_g(data, "company_name", default="[Company Name]"))
    fiscal_year  = _g(data, "fiscal_year", default="2025")
    sector       = _g(data, "sector", default="Industrial")
    decision     = _g(data, "recommendation", "lending_decision", default="REVIEW")
    fin          = _g(data, "financial_data", default={})
    ml           = _g(data, "ml_scores", default={})
    rec          = _g(data, "recommendation", default={})

    dec_color_hex = (
        "1B7A2B" if "APPROVE" in str(decision).upper() and "REJECT" not in str(decision).upper()
        else "C62828" if "REJECT" in str(decision).upper()
        else "E68A00"
    )
    dec_rgb = _risk_color(decision)

    # ── TOP HEADER BANNER (full-width navy table) ────────────────────────
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Left cell — firm name
    lc = header_tbl.rows[0].cells[0]
    _set_cell_shading(lc, "0A1F3C")
    lc.width = Cm(13)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run("YAKṢARĀJA")
    lr.bold = True
    lr.font.size = Pt(18)
    lr.font.color.rgb = WHITE
    lp2 = lc.add_paragraph()
    lr2 = lp2.add_run("AI Credit Decisioning Engine")
    lr2.font.size = Pt(9)
    lr2.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)
    lp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Right cell — date + confidential
    rc = header_tbl.rows[0].cells[1]
    _set_cell_shading(rc, "0A1F3C")
    rc.width = Cm(7)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(datetime.now().strftime("%d %B %Y"))
    rr.font.size = Pt(10)
    rr.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)
    rp2 = rc.add_paragraph()
    rr2 = rp2.add_run("STRICTLY CONFIDENTIAL")
    rr2.font.size = Pt(8)
    rr2.bold = True
    rr2.font.color.rgb = RGBColor(0xFF, 0xA0, 0x70)
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")

    # ── ORANGE ACCENT LINE ───────────────────────────────────────────────
    accent_tbl = doc.add_table(rows=1, cols=1)
    acc_cell = accent_tbl.rows[0].cells[0]
    _set_cell_shading(acc_cell, "E86C00")
    acc_cell.paragraphs[0].add_run(" ")
    for row in accent_tbl.rows:
        row.height = Cm(0.15)

    # ── MAIN TITLE ───────────────────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph("")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("CREDIT APPRAISAL MEMORANDUM")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = NAVY

    doc.add_paragraph("")

    # ── COMPANY NAME ─────────────────────────────────────────────────────
    comp_p = doc.add_paragraph()
    comp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = comp_p.add_run(company_name)
    cr.bold = True
    cr.font.size = Pt(22)
    cr.font.color.rgb = DARK_BLUE

    # ── SECTOR | FY ──────────────────────────────────────────────────────
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run(f"{sector}  ·  Fiscal Year {fiscal_year}")
    mr.font.size = Pt(12)
    mr.font.color.rgb = GREY

    for _ in range(2):
        doc.add_paragraph("")

    # ── RECOMMENDATION BADGE TABLE ───────────────────────────────────────
    badge_tbl = doc.add_table(rows=1, cols=1)
    badge_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    bc = badge_tbl.rows[0].cells[0]
    _set_cell_shading(bc, dec_color_hex)
    bc.width = Cm(10)
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run(f"RECOMMENDATION:  {str(decision).replace('_', ' ')}")
    br.bold = True
    br.font.size = Pt(16)
    br.font.color.rgb = WHITE

    for _ in range(2):
        doc.add_paragraph("")

    # ── KEY METRICS SUMMARY TABLE ─────────────────────────────────────────
    _add_para(doc, "KEY METRICS AT A GLANCE", bold=True, size=10,
              color=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    dscr_val    = _g(fin, "dscr", default=None)
    pd_val      = _g(ml,  "ensemble_pd", default=_g(fin, "ensemble_pd", default=None))
    icr_val     = _g(fin, "interest_coverage", default=None)
    de_val      = _g(fin, "debt_to_equity", default=None)
    rev_val     = _g(fin, "revenue", default=None)
    limit_val   = _g(rec, "recommended_limit_cr", default=None)
    rate_val    = _g(rec, "recommended_rate_pct", default=None)

    def _mv(v, fmt_str=".2f", prefix="", suffix=""):
        if v is None or v == "Not Given": return "N/A"
        try: return f"{prefix}{float(v):{fmt_str}}{suffix}"
        except: return str(v)

    metrics_rows = [
        ["Revenue (FY{})" .format(fiscal_year), _mv(rev_val, ",.0f", "₹", " Cr"),
         "Interest Coverage",                   _mv(icr_val, ".2f", "", "x")],
        ["DSCR",                                _mv(dscr_val, ".2f", "", "x"),
         "Debt / Equity",                       _mv(de_val, ".2f", "", "x")],
        ["Ensemble PD",                         _mv(pd_val, ".2%") if pd_val not in (None, "Not Given") else "N/A",
         "Recommended Rate",                    _mv(rate_val, ".2f", "", "%")],
        ["Credit Limit",                        _mv(limit_val, ",.1f", "₹", " Cr"),
         "Analysis Date",                       datetime.now().strftime("%d %b %Y")],
    ]

    metrics_tbl = doc.add_table(rows=len(metrics_rows), cols=4)
    metrics_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_data in enumerate(metrics_rows):
        for c_idx, val in enumerate(row_data):
            cell = metrics_tbl.rows[r_idx].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            is_label = (c_idx % 2 == 0)
            run = para.add_run(str(val))
            run.font.size = Pt(9)
            run.bold = is_label
            if is_label:
                run.font.color.rgb = NAVY
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_cell_shading(cell, "EEF2F8")
            else:
                run.font.color.rgb = BLACK
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if r_idx % 2 == 1:
                    _set_cell_shading(cell, "F9FAFB")

    for _ in range(3):
        doc.add_paragraph("")

    # ── FOOTER DISCLAIMER ────────────────────────────────────────────────
    footer_tbl = doc.add_table(rows=1, cols=1)
    fc = footer_tbl.rows[0].cells[0]
    _set_cell_shading(fc, "F0F0F0")
    fp = fc.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "This document is prepared for the exclusive use of Vivriti Capital's Credit Committee. "
        "Distribution outside the committee requires prior written authorization."
    )
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = GREY

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 2 — EXECUTIVE SUMMARY                                           •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_executive_summary(doc, data: dict):
    """Decision, limit, rate, and key points at a glance."""
    _add_styled_heading(doc, "1. Executive Summary", level=1)

    rec = _g(data, "recommendation", default={})
    ml = _g(data, "ml_scores", default={})
    fin = _g(data, "financial_data", default={})

    # Decision box
    decision = _g(rec, "lending_decision", default=_g(ml, "lending_decision", default="REVIEW"))
    limit = _g(rec, "recommended_limit_cr", default="Not Given")
    rate = _g(rec, "recommended_rate_pct", default="Not Given")

    _add_risk_badge(doc, "Lending Decision", decision)
    _add_key_value(doc, "Recommended Credit Limit", f"₹{_fmt(limit)} Cr" if limit not in ("N/A", "Not Given") else "Not Given")
    _add_key_value(doc, "Recommended Interest Rate", f"{_fmt(rate)}%" if rate not in ("N/A", "Not Given") else "Not Given")
    _add_key_value(doc, "Ensemble Probability of Default",
                   _pct(_g(ml, "ensemble_pd", default=_g(fin, "ensemble_pd"))))
    _add_key_value(doc, "DSCR", _fmt(_g(fin, "dscr", default="Not Given")))

    # Key conditions
    conditions = _g(rec, "key_conditions", default=[])
    if conditions and isinstance(conditions, list):
        _add_para(doc, "Key Conditions:", bold=True, size=10, space_after=2)
        for cond in conditions:
            _add_para(doc, f"  • {cond}", size=9, space_after=1)

    # Final rationale
    rationale = _g(rec, "final_rationale", default="")
    if rationale and rationale not in ("N/A", "Not Given"):
        doc.add_paragraph("")
        _add_para(doc, "Rationale:", bold=True, size=10, space_after=2)
        _add_para(doc, str(rationale), size=10, italic=True, space_after=6)

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 3 — COMPANY BACKGROUND                                          •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_company_background(doc, data: dict):
    """Sector, history, ownership structure."""
    _add_styled_heading(doc, "2. Company Background", level=1)

    fin = _g(data, "financial_data", default={})
    research = _g(data, "research", default={})

    _add_key_value(doc, "Company Name", _g(data, "company_name"))
    _add_key_value(doc, "Sector", _g(fin, "sector", default=_g(data, "sector", default="Not Given")))
    _add_key_value(doc, "Fiscal Year Under Review", str(_g(data, "fiscal_year")))
    _add_key_value(doc, "Promoter Holding", _pct(_g(fin, "promoter_holding_pct")))
    _add_key_value(doc, "Promoter Pledge", _pct(_g(fin, "promoter_pledge_pct")))
    _add_key_value(doc, "Institutional Holding", _pct(_g(fin, "institutional_holding_pct")))

    # Research summary
    news = _g(research, "company_news_summary", default="")
    if news and news not in ("N/A", "Not Given"):
        _add_styled_heading(doc, "Industry & Market Context", level=2)
        for line in str(news).split("\n"):
            _add_md_line(doc, line, heading_color=DARK_BLUE, text_color=BLACK, base_size=10)

    _add_key_value(doc, "Industry Outlook", _g(research, "industry_outlook", default="Not Given"),
                   color=_risk_color(_g(research, "industry_outlook", default="NEUTRAL")))

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 4 — FINANCIAL ANALYSIS                                           •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_financial_analysis(doc, data: dict):
    """5-year ratio tables and commentary."""
    _add_styled_heading(doc, "3. Financial Analysis", level=1)

    fin = _g(data, "financial_data", default={})

    # Key P&L metrics
    _add_styled_heading(doc, "Income Statement Highlights", level=2)
    pl_headers = ["Metric", "Value"]
    pl_rows = [
        ["Revenue", f"₹{_fmt(_g(fin, 'revenue'))} Cr"],
        ["EBITDA", f"₹{_fmt(_g(fin, 'ebitda'))} Cr"],
        ["EBITDA Margin", _pct(_g(fin, "ebitda_margin"))],
        ["PAT", f"₹{_fmt(_g(fin, 'pat'))} Cr"],
        ["Net Margin", _pct(_g(fin, "net_margin"))],
        ["Gross Margin", _pct(_g(fin, "gross_margin"))],
    ]
    _create_table(doc, pl_headers, pl_rows)

    doc.add_paragraph("")

    # Balance sheet
    _add_styled_heading(doc, "Balance Sheet Highlights", level=2)
    bs_headers = ["Metric", "Value"]
    bs_rows = [
        ["Total Assets", f"₹{_fmt(_g(fin, 'total_assets'))} Cr"],
        ["Total Equity", f"₹{_fmt(_g(fin, 'total_equity'))} Cr"],
        ["Total Debt", f"₹{_fmt(_g(fin, 'total_debt'))} Cr"],
        ["LT Borrowings", f"₹{_fmt(_g(fin, 'lt_borrowings'))} Cr"],
        ["ST Borrowings", f"₹{_fmt(_g(fin, 'st_borrowings'))} Cr"],
        ["Trade Receivables", f"₹{_fmt(_g(fin, 'trade_receivables'))} Cr"],
        ["Inventories", f"₹{_fmt(_g(fin, 'inventories'))} Cr"],
        ["Cash & Equivalents", f"₹{_fmt(_g(fin, 'cash_equivalents'))} Cr"],
    ]
    _create_table(doc, bs_headers, bs_rows)

    doc.add_paragraph("")

    # Key ratios
    _add_styled_heading(doc, "Key Financial Ratios", level=2)
    ratio_headers = ["Ratio", "Value", "Assessment"]
    dscr_val = _g(fin, "dscr", default=0)
    icr_val = _g(fin, "interest_coverage", default=0)
    de_val = _g(fin, "debt_to_equity", default=0)
    cr_val = _g(fin, "current_ratio", default=0)

    def _assess(val, good_thresh, warn_thresh, higher_is_better=True):
        try:
            v = float(val)
        except (ValueError, TypeError):
            return "N/A"
        if higher_is_better:
            if v >= good_thresh: return "GREEN"
            elif v >= warn_thresh: return "AMBER"
            else: return "RED"
        else:
            if v <= good_thresh: return "GREEN"
            elif v <= warn_thresh: return "AMBER"
            else: return "RED"

    ratio_rows = [
        ["DSCR", _fmt(dscr_val), _assess(dscr_val, 1.5, 1.0)],
        ["Interest Coverage Ratio", _fmt(icr_val), _assess(icr_val, 2.0, 1.3)],
        ["Debt-to-Equity", _fmt(de_val), _assess(de_val, 1.5, 2.5, higher_is_better=False)],
        ["Current Ratio", _fmt(cr_val), _assess(cr_val, 1.2, 0.8)],
        ["ROE", _pct(_g(fin, "roe")), _assess(_g(fin, "roe", default=0), 0.12, 0.05)],
        ["ROA", _pct(_g(fin, "roa")), _assess(_g(fin, "roa", default=0), 0.05, 0.02)],
    ]
    _create_table(doc, ratio_headers, ratio_rows)

    doc.add_paragraph("")

    # Cash flow
    _add_styled_heading(doc, "Cash Flow Analysis", level=2)
    cf_headers = ["Metric", "Value"]
    cf_rows = [
        ["Cash from Operations", f"₹{_fmt(_g(fin, 'cfo'))} Cr"],
        ["Cash from Investing", f"₹{_fmt(_g(fin, 'cfi'))} Cr"],
        ["Cash from Financing", f"₹{_fmt(_g(fin, 'cff'))} Cr"],
        ["Capital Expenditure", f"₹{_fmt(_g(fin, 'capex'))} Cr"],
        ["Free Cash Flow", f"₹{_fmt(_g(fin, 'free_cash_flow'))} Cr"],
    ]
    _create_table(doc, cf_headers, cf_rows)

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 5 — FINANCIAL FORENSICS                                          •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_forensics_section(doc, data: dict):
    """Beneish M-Score, Altman Z-Score, Piotroski F-Score, audit flags."""
    _add_styled_heading(doc, "4. Financial Forensics", level=1)

    forensics = _g(data, "forensics", default={})
    fin = _g(data, "financial_data", default={})

    _add_para(doc, "Automated forensic analysis using three established academic models "
              "to detect potential earnings manipulation, bankruptcy risk, and financial strength.",
              size=10, italic=True)

    # Beneish M-Score
    _add_styled_heading(doc, "Beneish M-Score (Earnings Manipulation)", level=2)
    m_score = _g(forensics, "beneish_m_score", default=_g(fin, "beneish_m_score", default="Not Given"))
    m_flag = "SUSPICIOUS" if m_score not in ("N/A", "Not Given") and float(m_score) > -2.22 else "CLEAN"
    _add_risk_badge(doc, f"M-Score: {_fmt(m_score)}", m_flag)
    _add_para(doc, "Threshold: > -2.22 indicates potential manipulation (Beneish 1999).", size=9, color=GREY)

    # Altman Z-Score
    _add_styled_heading(doc, "Altman Z-Score (Bankruptcy Risk)", level=2)
    z_score = _g(forensics, "altman_z_score", default=_g(fin, "altman_z_score", default="Not Given"))
    if z_score not in ("N/A", "Not Given"):
        z_val = float(z_score)
        z_flag = "SAFE" if z_val > 2.99 else ("GREY ZONE" if z_val > 1.81 else "DISTRESS")
    else:
        z_flag = "N/A"
    _add_risk_badge(doc, f"Z-Score: {_fmt(z_score)}", z_flag)
    _add_para(doc, "Zones: >2.99 = Safe, 1.81-2.99 = Grey, <1.81 = Distress (Altman 1968).", size=9, color=GREY)

    # Piotroski F-Score
    _add_styled_heading(doc, "Piotroski F-Score (Financial Strength)", level=2)
    f_score = _g(forensics, "piotroski_f_score", default=_g(fin, "piotroski_f_score", default="Not Given"))
    if f_score not in ("N/A", "Not Given"):
        f_val = int(float(f_score))
        f_flag = "STRONG" if f_val >= 7 else ("MODERATE" if f_val >= 4 else "WEAK")
    else:
        f_flag = "N/A"
    _add_risk_badge(doc, f"F-Score: {f_score}/9", f_flag)

    # Audit quality flags
    _add_styled_heading(doc, "Audit Quality Indicators", level=2)
    audit_headers = ["Indicator", "Value", "Flag"]
    audit_rows = [
        ["Auditor Distress Score", f"{_g(fin, 'auditor_distress_score', default='N/A')}/5",
         "RED" if _g(fin, "auditor_distress_score", default=0) >= 3 else "GREEN"],
        ["Going Concern Opinion", "Yes" if _g(fin, "going_concern_flag", default=0) else "No",
         "RED" if _g(fin, "going_concern_flag", default=0) else "GREEN"],
        ["Qualified Opinion", "Yes" if _g(fin, "qualified_opinion_flag", default=0) else "No",
         "RED" if _g(fin, "qualified_opinion_flag", default=0) else "GREEN"],
        ["Auditor Resigned", "Yes" if _g(fin, "auditor_resigned_flag", default=0) else "No",
         "RED" if _g(fin, "auditor_resigned_flag", default=0) else "GREEN"],
        ["Related Party Tx/Revenue", _pct(_g(fin, "related_party_tx_to_rev")),
         "AMBER" if _g(fin, "related_party_tx_to_rev", default=0) > 0.10 else "GREEN"],
    ]
    _create_table(doc, audit_headers, audit_rows)

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 6 — PROMOTER NETWORK ANALYSIS                                    •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_network_section(doc, data: dict):
    """Promoter network description and contagion risk."""
    _add_styled_heading(doc, "5. Promoter Network & Contagion Risk", level=1)

    network = _g(data, "network", default={})
    fin = _g(data, "financial_data", default={})

    _add_para(doc, "Analysis of the promoter's director network using MCA21 data to identify "
              "contagion risk from related entities.", size=10, italic=True)

    contagion = _g(network, "contagion_risk_score",
                   default=_g(fin, "contagion_risk_score"))
    if not _not_given(contagion):
        try:
            level = "HIGH" if float(contagion) > 0.5 else ("MEDIUM" if float(contagion) > 0.25 else "LOW")
        except (ValueError, TypeError):
            level = "Not Given"
    else:
        level = "Not Given"
    _add_risk_badge(doc, f"Contagion Risk Score: {_fmt(contagion)}", level)

    net_headers = ["Indicator", "Value"]
    net_rows = [
        ["Promoter Total Companies", str(_g(fin, "promoter_total_companies"))],
        ["Promoter NPA Companies", str(_g(fin, "promoter_npa_companies"))],
        ["Promoter Struck-Off Companies", str(_g(fin, "promoter_struck_off_companies"))],
        ["DIN Disqualified Directors", str(_g(fin, "din_disqualified_count"))],
        ["Network NPA Ratio", _pct(_g(fin, "network_npa_ratio"))],
        ["Customer Concentration", _pct(_g(fin, "customer_concentration"))],
        ["Supplier Concentration", _pct(_g(fin, "supplier_concentration"))],
    ]
    _create_table(doc, net_headers, net_rows)

    # Network insights if available
    insights = _g(network, "insights", default="")
    if insights and insights not in ("N/A", "Not Given"):
        doc.add_paragraph("")
        _add_para(doc, "Network Insights:", bold=True, size=10)
        _add_para(doc, str(insights), size=10)

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 7 — SATELLITE & OPERATIONAL REALITY                              •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_satellite_section(doc, data: dict):
    """Satellite-based operational activity assessment."""
    _add_styled_heading(doc, "6. Satellite Activity & Operational Reality", level=1)

    satellite = _g(data, "satellite", default={})
    fin = _g(data, "financial_data", default={})

    _add_para(doc, "Independent verification of operational activity using Sentinel-2 satellite "
              "imagery analysis of company premises.", size=10, italic=True)

    score = _g(satellite, "activity_score",
               default=_g(fin, "satellite_activity_score", default="Not Given"))
    category = _g(satellite, "activity_category",
                  default=_g(fin, "satellite_activity_category", default="Not Given"))
    rev_flag = _g(satellite, "vs_revenue_flag",
                  default=_g(fin, "satellite_vs_revenue_flag", default=0))

    _add_risk_badge(doc, f"Activity Score: {_fmt(score, '.1f')}/100", str(category))

    sat_headers = ["Indicator", "Value", "Assessment"]
    sat_rows = [
        ["Activity Score", _fmt(score, ".1f"), str(category)],
        ["Activity vs Revenue Mismatch", "Yes" if rev_flag else "No",
         "RED" if rev_flag else "GREEN"],
    ]
    _create_table(doc, sat_headers, sat_rows)

    # GST cross-verification
    gst = _g(data, "gst", default={})
    _add_styled_heading(doc, "GST Cross-Verification", level=2)
    gst_headers = ["Indicator", "Value", "Flag"]
    gst_rows = [
        ["GST-Bank Divergence", _pct(_g(fin, "gst_vs_bank_divergence")),
         "RED" if _g(fin, "gst_divergence_flag", default=0) else "GREEN"],
        ["GST Filing Delays", str(_g(fin, "gst_filing_delays_count", default="Not Given")),
         "AMBER" if _g(fin, "gst_filing_delays_count", default=0) > 2 else "GREEN"],
        ["E-Way Bill Consistency", _pct(_g(fin, "ewaybill_volume_consistency")),
         "GREEN" if _g(fin, "ewaybill_volume_consistency", default=1) > 0.8 else "AMBER"],
        ["GST Payment Delay (days)", _fmt(_g(fin, "gst_payment_delay_days"), ".0f"),
         "RED" if _g(fin, "gst_payment_delay_days", default=0) > 30 else "GREEN"],
    ]
    _create_table(doc, gst_headers, gst_rows)

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 8 — STRESS TESTING                                               •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_stress_test_section(doc, data: dict):
    """Monte Carlo stress test results — P10/P50/P90 and named scenarios."""
    _add_styled_heading(doc, "7. Stress Test Results", level=1)

    stress = _g(data, "stress_test", default={})

    _add_para(doc, "Monte Carlo simulation (10,000 paths) projecting DSCR under "
              "various macro-economic scenarios.", size=10, italic=True)

    # Percentile table
    _add_styled_heading(doc, "DSCR Distribution (Simulated)", level=2)
    p10 = _g(stress, "dscr_p10", default="Not Given")
    p50 = _g(stress, "dscr_p50", default="Not Given")
    p90 = _g(stress, "dscr_p90", default="Not Given")
    breach_prob = _g(stress, "covenant_breach_probability", default="Not Given")

    pct_headers = ["Percentile", "Simulated DSCR", "Assessment"]
    pct_rows = [
        ["P10 (Severe Stress)", _fmt(p10),
         "RED" if p10 not in ("N/A", "Not Given") and float(p10) < 1.0 else "AMBER"],
        ["P50 (Base Case)", _fmt(p50),
         "GREEN" if p50 not in ("N/A", "Not Given") and float(p50) > 1.2 else "AMBER"],
        ["P90 (Optimistic)", _fmt(p90),
         "GREEN" if p90 not in ("N/A", "Not Given") and float(p90) > 1.5 else "GREEN"],
    ]
    _create_table(doc, pct_headers, pct_rows)

    if breach_prob not in ("N/A", "Not Given"):
        doc.add_paragraph("")
        _add_risk_badge(doc, "Probability of Covenant Breach (DSCR < 1.0)", _pct(breach_prob))

    # Named scenarios
    scenarios = _g(stress, "named_scenarios", default=[])
    if scenarios and isinstance(scenarios, list):
        _add_styled_heading(doc, "Named Stress Scenarios", level=2)
        sc_headers = ["Scenario", "DSCR Impact", "PD Impact"]
        sc_rows = []
        for sc in scenarios:
            if isinstance(sc, dict):
                sc_rows.append([
                    _g(sc, "name", default="Scenario"),
                    _fmt(_g(sc, "dscr_impact")),
                    _pct(_g(sc, "pd_impact")),
                ])
        if sc_rows:
            _create_table(doc, sc_headers, sc_rows)

    # DNA matching
    dna = _g(data, "dna_match", default={})
    closest = _g(dna, "closest_default_archetype", default="Not Given")
    similarity = _g(dna, "max_archetype_similarity", default="Not Given")
    if closest not in ("N/A", "Not Given"):
        doc.add_paragraph("")
        _add_styled_heading(doc, "Default DNA Matching", level=2)
        _add_key_value(doc, "Closest Default Archetype", str(closest))
        _add_key_value(doc, "Similarity Score", _fmt(similarity))

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 9 — MANAGEMENT QUALITY (CEO INTERVIEW)                          •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_management_section(doc, data: dict):
    """CEO interview insights and management quality score.
    Rule 5: Only show detailed scores if transcript was provided."""
    _add_styled_heading(doc, "8. Management Quality Assessment", level=1)

    ceo = _g(data, "ceo_interview", default={})
    scores = _g(ceo, "key_scores", default={})
    used_fallback = _g(ceo, "used_fallback", default=True)

    # Rule 5: Gate on transcript availability
    if used_fallback:
        _add_para(doc, "Management Quality Assessment: INSUFFICIENT DATA",
                  size=12, bold=True, color=AMBER)
        _add_para(doc, "No earnings call transcript or interview audio was provided. "
                  "Management quality scores shown below are proxy estimates "
                  "derived from financial data and should not be relied upon for "
                  "the credit decision. Schedule management interaction before "
                  "final approval.",
                  size=10, italic=True, color=GREY)
        doc.add_paragraph("")

    _add_para(doc, "Assessment based on structured CEO/promoter interview analysis — scoring "
              "consistency between financial data and management narrative.",
              size=10, italic=True)

    # MQ score
    mq = _g(ceo, "management_quality_score", default="Not Given")
    if mq not in ("N/A", "Not Given"):
        level = "GREEN" if float(mq) >= 65 else ("AMBER" if float(mq) >= 40 else "RED")
        _add_risk_badge(doc, f"Management Quality Score: {_fmt(mq, '.1f')}/100", level)

    # Detailed scores
    score_headers = ["Metric", "Score", "Interpretation"]
    score_rows = [
        ["Overall Sentiment", _fmt(_g(scores, "ceo_sentiment_overall")),
         "Balanced" if 0.2 <= _g(scores, "ceo_sentiment_overall", default=0.5) <= 0.6 else "Review"],
        ["Debt Discussion Sentiment", _fmt(_g(scores, "ceo_sentiment_debt")),
         "Honest" if _g(scores, "ceo_sentiment_debt", default=0) < 0.3 else "Suspicious"],
        ["Deflection Score", _fmt(_g(scores, "ceo_deflection_score")),
         "HIGH" if _g(scores, "ceo_deflection_score", default=0) > 0.4 else "OK"],
        ["Overconfidence Score", _fmt(_g(scores, "ceo_overconfidence_score")),
         "HIGH" if _g(scores, "ceo_overconfidence_score", default=0) > 0.3 else "OK"],
        ["Specificity Score", _fmt(_g(scores, "ceo_specificity_score")),
         "Good" if _g(scores, "ceo_specificity_score", default=0) > 0.3 else "Low"],
    ]
    _create_table(doc, score_headers, score_rows)

    # Red flags
    flags = _g(ceo, "red_flags", default=[])
    if flags and isinstance(flags, list):
        doc.add_paragraph("")
        _add_para(doc, "Interview Red Flags:", bold=True, size=10, color=RED)
        for flag in flags:
            if isinstance(flag, dict):
                _add_para(doc, f"  🚩 [{_g(flag, 'severity')}] {_g(flag, 'description')}",
                          size=9, color=RED)

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 10 — BULL VS BEAR (ADVERSARIAL DEBATE)                          •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_bull_bear_section(doc, data: dict):
    """Side-by-side adversarial bull and bear case arguments."""
    _add_styled_heading(doc, "9. Adversarial Credit Committee Debate", level=1)

    _add_para(doc, "Two independent LLM agents debated this application — one building "
              "the strongest case for approval, the other seeking every reason to reject. "
              "This mirrors real credit committee dynamics with a devil's advocate.",
              size=10, italic=True)

    # Bull case
    _add_styled_heading(doc, "🟢 Bull Case (Approval Agent)", level=2)
    bull = _g(data, "bull_case", default="Bull case not available.")
    for line in str(bull).split("\n"):
        _add_md_line(doc, line, heading_color=GREEN, text_color=None, base_size=10)

    _add_page_break(doc)

    # Bear case
    _add_styled_heading(doc, "🔴 Bear Case (Dissent Agent)", level=2)
    bear = _g(data, "bear_case", default="Bear case not available.")
    for line in str(bear).split("\n"):
        _add_md_line(doc, line, heading_color=RED, text_color=None, base_size=10)

    _add_page_break(doc)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  SECTION 11 — FINAL RECOMMENDATION                                        •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_recommendation(doc, data: dict):
    """Final verdict with conditions and signatures."""
    _add_styled_heading(doc, "10. Final Recommendation & Decision", level=1)

    rec = _g(data, "recommendation", default={})

    # Decision
    decision = _g(rec, "lending_decision", default="REVIEW")
    _add_risk_badge(doc, "FINAL DECISION", decision)

    doc.add_paragraph("")

    # Terms
    _add_key_value(doc, "Recommended Credit Limit",
                   f"₹{_fmt(_g(rec, 'recommended_limit_cr'))} Cr")
    _add_key_value(doc, "Recommended Interest Rate",
                   f"{_fmt(_g(rec, 'recommended_rate_pct'))}%")

    # Bull & Bear summaries
    doc.add_paragraph("")
    _add_para(doc, "Arguments For (Bull Summary):", bold=True, size=10, color=GREEN)
    _add_para(doc, str(_g(rec, "bull_summary", default="Not Given")), size=10, space_after=6)

    _add_para(doc, "Arguments Against (Bear Summary):", bold=True, size=10, color=RED)
    _add_para(doc, str(_g(rec, "bear_summary", default="Not Given")), size=10, space_after=6)

    # Key conditions
    conditions = _g(rec, "key_conditions", default=[])
    if conditions and isinstance(conditions, list):
        _add_para(doc, "Key Conditions & Covenants:", bold=True, size=11)
        for i, cond in enumerate(conditions, 1):
            _add_para(doc, f"  {i}. {cond}", size=10, space_after=2)

    # Final rationale
    doc.add_paragraph("")
    _add_para(doc, "Final Rationale:", bold=True, size=11)
    _add_para(doc, str(_g(rec, "final_rationale", default="Not Given")), size=10, italic=True)

    # Signature block
    doc.add_paragraph("")
    doc.add_paragraph("")
    _add_para(doc, "─" * 50, size=8, color=GREY)

    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["Credit Analyst", "Risk Officer", "Credit Committee Head"]
    for i, label in enumerate(labels):
        sig_table.rows[0].cells[i].text = label
        for para in sig_table.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        sig_table.rows[1].cells[i].text = "\n\nSignature: ______________\nDate: __/__/____"
        for para in sig_table.rows[1].cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(8)

    doc.add_paragraph("")
    _add_para(doc, "— End of Credit Appraisal Memorandum —",
              size=10, italic=True, color=GREY, alignment=WD_ALIGN_PARAGRAPH.CENTER)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  PART C — DOCX ASSEMBLY                                                   •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def _setup_document() -> Any:
    """Create a new Document with professional styling and page setup."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = BLACK

    # Heading 1 style
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.color.rgb = NAVY
    h1.font.bold = True

    # Heading 2 style
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.color.rgb = DARK_BLUE
    h2.font.bold = True

    return doc


def _add_header_footer(doc, company_name: str):
    """Add headers and footers with page numbers and confidentiality."""
    for section in doc.sections:
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.add_run(f"CONFIDENTIAL — {company_name} — Credit Appraisal Memo")
        run.font.size = Pt(7)
        run.font.color.rgb = GREY
        run.italic = True

        # Footer with page numbers
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run("Yakṣarāja | Vivriti Capital | Page ")
        run.font.size = Pt(7)
        run.font.color.rgb = GREY
        # Page number field
        fld_xml = (
            '<w:fldSimple {} w:instr=" PAGE "/>'.format(nsdecls("w"))
        )
        fld = parse_xml(fld_xml)
        footer_para._p.append(fld)


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  PART D — SAVE OUTPUT (DOCX + JSON)                                       •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def _save_summary_json(data: dict, output_dir: str) -> str:
    """Save a summary scores JSON alongside the DOCX."""
    company_name = _g(data, "company_name", default="unknown")
    safe_name = str(company_name).replace(" ", "_").replace("/", "_")
    json_path = os.path.join(output_dir, f"scores_{safe_name}.json")

    rec = _g(data, "recommendation", default={})
    ml = _g(data, "ml_scores", default={})
    fin = _g(data, "financial_data", default={})
    ceo = _g(data, "ceo_interview", default={})

    summary = {
        "company_name": str(company_name),
        "fiscal_year": _g(data, "fiscal_year"),
        "generated_at": datetime.now().isoformat(),
        "lending_decision": _g(rec, "lending_decision", default="Not Given"),
        "recommended_limit_cr": _g(rec, "recommended_limit_cr", default="Not Given"),
        "recommended_rate_pct": _g(rec, "recommended_rate_pct", default="Not Given"),
        "ensemble_pd": _g(ml, "ensemble_pd", default=_g(fin, "ensemble_pd")),
        "dscr": _g(fin, "dscr"),
        "interest_coverage": _g(fin, "interest_coverage"),
        "debt_to_equity": _g(fin, "debt_to_equity"),
        "beneish_m_score": _g(fin, "beneish_m_score"),
        "contagion_risk_score": _g(fin, "contagion_risk_score"),
        "satellite_activity_score": _g(fin, "satellite_activity_score"),
        "management_quality_score": _g(ceo, "management_quality_score"),
        "key_conditions": _g(rec, "key_conditions", default=[]),
    }

    try:
        os.makedirs(output_dir, exist_ok=True)
        with open(json_path, "w") as f:
            json.dump(summary, f, indent=2, default=str)
        logger.info(f"Summary JSON saved: {json_path}")
    except Exception as e:
        logger.error(f"Failed to save summary JSON: {e}")
        json_path = ""

    return json_path


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  MAIN ENTRY POINT                                                         •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

def generate_cam(
    all_data: dict,
    output_dir: str = "data/processed/",
) -> str:
    """
    Generate a complete Credit Appraisal Memo (CAM) as a DOCX file.

    Assembles all outputs from Person 1 + Person 2 + Person 3 into a
    professional ~20-page document.

    Args:
        all_data: Pipeline output dict containing:
            - company_name, fiscal_year
            - financial_data (dict/Series from Person 1)
            - forensics, ml_scores, trajectory (Person 1)
            - network, stress_test, dna_match, satellite, gst (Person 2)
            - research, ceo_interview, bull_case, bear_case, recommendation (Person 3)
        output_dir: Directory to save output files (default: data/processed/)

    Returns:
        Absolute path to the generated DOCX file.
    """
    if not DOCX_AVAILABLE:
        logger.error("python-docx is required for CAM generation. pip install python-docx")
        return ""

    company_name = _g(all_data, "company_name", default="Unknown_Company")
    safe_name = str(company_name).replace(" ", "_").replace("/", "_")
    date_str = datetime.now().strftime("%Y%m%d")

    logger.info(f"{'='*60}")
    logger.info(f"GENERATING CAM: {company_name}")
    logger.info(f"{'='*60}")

    # Create output directory
    os.makedirs(output_dir, exist_ok=True)

    # ”—€ Build the DOCX ”——————————————————————————————————————————————————€
    doc = _setup_document()

    # Convert financial_data from pd.Series to dict if needed
    fin_data = all_data.get("financial_data", {})
    if hasattr(fin_data, "to_dict"):
        all_data["financial_data"] = fin_data.to_dict()

    # ”—€ Run consistency validation (Rules 4, 5, 6, 7) ”———————————————————€
    logger.info("Running CAM consistency validation...")
    validation_issues = validate_cam_consistency(all_data)
    if validation_issues:
        logger.warning(f"CAM validation found {len(validation_issues)} issue(s)")

    logger.info("Generating Section 1: Cover Page")
    generate_cover_page(doc, all_data)

    logger.info("Generating Section 2: Executive Summary")
    generate_executive_summary(doc, all_data)

    logger.info("Generating Section 3: Company Background")
    generate_company_background(doc, all_data)

    logger.info("Generating Section 4: Financial Analysis")
    generate_financial_analysis(doc, all_data)

    logger.info("Generating Section 5: Financial Forensics")
    generate_forensics_section(doc, all_data)

    logger.info("Generating Section 6: Network Analysis")
    generate_network_section(doc, all_data)

    logger.info("Generating Section 7: Satellite & GST")
    generate_satellite_section(doc, all_data)

    logger.info("Generating Section 8: Stress Testing")
    generate_stress_test_section(doc, all_data)

    logger.info("Generating Section 9: Management Quality")
    generate_management_section(doc, all_data)

    logger.info("Generating Section 10: Bull vs Bear Debate")
    generate_bull_bear_section(doc, all_data)

    logger.info("Generating Section 11: Final Recommendation")
    generate_recommendation(doc, all_data)

    # Add headers/footers
    _add_header_footer(doc, str(company_name))

    # ”—€ Save DOCX ”———————————————————————————————————————————————————————€
    docx_filename = f"CAM_{safe_name}_{date_str}.docx"
    docx_path = os.path.join(output_dir, docx_filename)

    try:
        doc.save(docx_path)
        docx_path = os.path.abspath(docx_path)
        logger.info(f"œ… CAM saved: {docx_path}")
    except Exception as e:
        logger.error(f"Failed to save DOCX: {e}")
        return ""

    # ”—€ Save summary JSON ”———————————————————————————————————————————————€
    json_path = _save_summary_json(all_data, output_dir)

    logger.info(f"{'='*60}")
    logger.info(f"CAM GENERATION COMPLETE")
    logger.info(f"  DOCX: {docx_path}")
    logger.info(f"  JSON: {json_path}")
    logger.info(f"{'='*60}")

    return docx_path


# •”•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••—
# •‘  CLI — STANDALONE TEST                                                    •‘
# •š•••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••••

if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("CAM GENERATOR — Standalone Test (Demo Data)")
    print("=" * 60)

    demo_data = {
        "company_name": "Sunrise Textile Mills",
        "fiscal_year": 2024,
        "financial_data": {
            "sector": "Textiles",
            "revenue": 850.0, "ebitda": 127.5, "ebitda_margin": 0.15,
            "pat": 51.0, "net_margin": 0.06, "gross_margin": 0.38,
            "total_assets": 1200.0, "total_equity": 325.0, "total_debt": 520.0,
            "lt_borrowings": 350.0, "st_borrowings": 170.0,
            "trade_receivables": 120.0, "inventories": 95.0, "cash_equivalents": 45.0,
            "cfo": 95.0, "cfi": -55.0, "cff": -30.0, "capex": 55.0, "free_cash_flow": 42.0,
            "dscr": 1.85, "interest_coverage": 2.4, "debt_to_equity": 1.6,
            "current_ratio": 1.25, "roe": 0.14, "roa": 0.06,
            "promoter_holding_pct": None, "promoter_pledge_pct": None,
            "institutional_holding_pct": None,
            "beneish_m_score": -2.45, "altman_z_score": 2.3, "piotroski_f_score": 6,
            "auditor_distress_score": 1, "going_concern_flag": 0,
            "qualified_opinion_flag": 0, "auditor_resigned_flag": 0,
            "related_party_tx_to_rev": 0.05,
            "promoter_total_companies": None, "promoter_npa_companies": None,
            "promoter_struck_off_companies": None, "din_disqualified_count": None,
            "network_npa_ratio": None, "contagion_risk_score": None,
            "customer_concentration": 0.35, "supplier_concentration": 0.42,
            "satellite_activity_score": 82.5, "satellite_activity_category": "ACTIVE",
            "satellite_vs_revenue_flag": 0,
            "gst_vs_bank_divergence": 0.03, "gst_divergence_flag": 0,
            "gst_filing_delays_count": 1, "ewaybill_volume_consistency": 0.91,
            "gst_payment_delay_days": 12,
            "ensemble_pd": 0.12,
        },
        "forensics": {
            "beneish_m_score": -2.45, "altman_z_score": 2.3, "piotroski_f_score": 6,
        },
        "ml_scores": {
            "ensemble_pd": 0.12, "xgb_pd": 0.11, "rf_pd": 0.14, "lgb_pd": 0.13,
            "lending_decision": "APPROVE", "risk_premium": 3.5,
        },
        "trajectory": {"dscr_trend": "STABLE", "months_to_danger": 36},
        "network": {"contagion_risk_score": 0.15},
        "stress_test": {
            "dscr_p10": 1.05, "dscr_p50": 1.65, "dscr_p90": 2.15,
            "covenant_breach_probability": 0.08,
            "named_scenarios": [
                {"name": "Revenue -20%", "dscr_impact": 1.15, "pd_impact": 0.25},
                {"name": "Rate +200bps", "dscr_impact": 1.45, "pd_impact": 0.18},
                {"name": "Combined Shock", "dscr_impact": 0.85, "pd_impact": 0.42},
            ],
        },
        "dna_match": {"closest_default_archetype": "None (Healthy)", "max_archetype_similarity": 0.18},
        "satellite": {"activity_score": 82.5, "activity_category": "ACTIVE", "vs_revenue_flag": 0},
        "gst": {},
        "research": {
            "company_news_summary": "Sunrise Textile Mills reported 12% revenue growth in FY2024.",
            "industry_outlook": "POSITIVE",
            "research_sentiment_score": 0.72,
            "key_positives_found": ["PLI scheme support", "China+1 trend"],
            "key_risks_found": ["Raw material volatility"],
        },
        "ceo_interview": {
            "key_scores": {
                "ceo_sentiment_overall": 0.45,
                "ceo_sentiment_debt": 0.15,
                "ceo_deflection_score": 0.18,
                "ceo_overconfidence_score": 0.12,
                "ceo_specificity_score": 0.55,
            },
            "red_flags": [],
            "red_flag_count": 0,
            "management_quality_score": 72.5,
            "used_fallback": False,
        },
        "bull_case": (
            "## 1. EXECUTIVE SUMMARY\n"
            "Sunrise Textile Mills presents a strong lending opportunity with DSCR of 1.85x "
            "and 15% EBITDA margin.\n\n"
            "## 2. FINANCIAL STRENGTHS\n"
            "- Healthy DSCR of 1.85x\n- Interest coverage of 2.4x\n- Positive free cash flow ₹42 Cr\n"
        ),
        "bear_case": (
            "## 1. CRITICAL CONCERNS\n"
            "- D/E of 1.6x leaves limited buffer\n"
            "- Raw material price volatility could compress margins\n\n"
            "## 2. STRESS SCENARIO\n"
            "- Under 20% revenue shock, DSCR drops to ~1.20x\n"
        ),
        "recommendation": {
            "lending_decision": "CONDITIONAL_APPROVE",
            "recommended_limit_cr": 187.0,
            "recommended_rate_pct": 10.0,
            "key_conditions": [
                "DSCR floor covenant at 1.20x — quarterly monitoring",
                "Promoter personal guarantee covering 50% of exposure",
                "Quarterly GST cross-verification with bank statements",
                "Annual credit review with fresh financials",
            ],
            "bull_summary": "Strong DSCR, positive cash flow, and supportive industry dynamics.",
            "bear_summary": "Leverage and raw material risks require covenant protection.",
            "final_rationale": (
                "After weighing both perspectives, the committee recommends a CONDITIONAL APPROVE "
                "with ₹187 Cr limit at 10.0% interest rate. DSCR of 1.85x and ensemble PD of 12% "
                "support lending with appropriate risk mitigants."
            ),
        },
    }

    path = generate_cam(demo_data, output_dir="data/processed/")
    if path:
        print(f"\nœ… Demo CAM generated: {path}")
    else:
        print("\nŒ CAM generation failed")
    print("=" * 60)


